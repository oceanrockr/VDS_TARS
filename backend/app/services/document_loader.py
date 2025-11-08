"""
T.A.R.S. Document Loader Service
Multi-format document loading and chunking
Phase 3 - Document Indexing & RAG
"""

import hashlib
import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import chardet

# Document processing libraries
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
import pandas as pd

from ..core.config import settings
from ..models.rag import Document, DocumentMetadata, DocumentChunk

logger = logging.getLogger(__name__)


class DocumentLoaderService:
    """
    Service for loading and processing documents from multiple formats.

    Supported Formats:
    - PDF (.pdf)
    - Microsoft Word (.docx)
    - Plain Text (.txt)
    - Markdown (.md)
    - CSV (.csv)

    Features:
    - Automatic encoding detection
    - Text chunking with overlap
    - Metadata extraction
    - SHA256 hashing for deduplication
    - File size validation
    """

    def __init__(self):
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
        self.max_file_size = settings.MAX_FILE_SIZE_MB * 1024 * 1024  # Convert to bytes
        self.allowed_extensions = [
            ext.strip() for ext in settings.ALLOWED_EXTENSIONS.split(",")
        ]
        logger.info(f"DocumentLoaderService initialized")
        logger.info(f"Allowed extensions: {self.allowed_extensions}")

    def validate_file(self, file_path: str) -> Tuple[bool, Optional[str]]:
        """
        Validate file before processing.

        Args:
            file_path: Path to the file

        Returns:
            Tuple of (is_valid, error_message)
        """
        path = Path(file_path)

        # Check if file exists
        if not path.exists():
            return False, f"File not found: {file_path}"

        # Check if it's a file (not directory)
        if not path.is_file():
            return False, f"Not a file: {file_path}"

        # Check file extension
        ext = path.suffix.lower()
        if ext not in self.allowed_extensions:
            return False, f"Unsupported file type: {ext}"

        # Check file size
        file_size = path.stat().st_size
        if file_size > self.max_file_size:
            max_mb = self.max_file_size / (1024 * 1024)
            actual_mb = file_size / (1024 * 1024)
            return False, f"File too large: {actual_mb:.1f}MB (max {max_mb}MB)"

        if file_size == 0:
            return False, "File is empty"

        return True, None

    def compute_file_hash(self, file_path: str) -> str:
        """
        Compute SHA256 hash of file content.

        Args:
            file_path: Path to the file

        Returns:
            Hexadecimal hash string
        """
        sha256 = hashlib.sha256()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b''):
                sha256.update(chunk)
        return sha256.hexdigest()

    def detect_encoding(self, file_path: str) -> str:
        """
        Detect file encoding using chardet.

        Args:
            file_path: Path to the file

        Returns:
            Detected encoding (default: utf-8)
        """
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)  # Read first 10KB
                result = chardet.detect(raw_data)
                encoding = result['encoding']
                confidence = result['confidence']

                logger.debug(
                    f"Detected encoding: {encoding} "
                    f"(confidence: {confidence:.2f})"
                )

                return encoding if encoding else 'utf-8'

        except Exception as e:
            logger.warning(f"Encoding detection failed: {e}, using utf-8")
            return 'utf-8'

    def load_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Load PDF file and extract text.

        Args:
            file_path: Path to PDF file

        Returns:
            Tuple of (text_content, metadata)
        """
        text = ""
        metadata = {"page_count": 0, "extraction_method": "pdfplumber"}

        try:
            # Try pdfplumber first (better text extraction)
            with pdfplumber.open(file_path) as pdf:
                metadata["page_count"] = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"

            # Fallback to PyPDF2 if pdfplumber didn't extract text
            if not text.strip():
                logger.warning(f"pdfplumber failed, trying PyPDF2: {file_path}")
                metadata["extraction_method"] = "pypdf2"

                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    metadata["page_count"] = len(pdf_reader.pages)

                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"

            if not text.strip():
                logger.warning(
                    f"No text extracted from PDF (may be scanned): {file_path}"
                )
                if settings.ENABLE_OCR:
                    # TODO: Implement OCR fallback in future phase
                    pass

        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            raise

        return text.strip(), metadata

    def load_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Load Microsoft Word document.

        Args:
            file_path: Path to DOCX file

        Returns:
            Tuple of (text_content, metadata)
        """
        try:
            doc = DocxDocument(file_path)

            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)

            metadata = {
                "paragraph_count": len(paragraphs),
                "extraction_method": "python-docx"
            }

            return text.strip(), metadata

        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            raise

    def load_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Load plain text or markdown file.

        Args:
            file_path: Path to text file

        Returns:
            Tuple of (text_content, metadata)
        """
        try:
            encoding = self.detect_encoding(file_path)

            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()

            metadata = {
                "encoding": encoding,
                "line_count": text.count('\n') + 1
            }

            return text.strip(), metadata

        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {e}")
            raise

    def load_csv(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Load CSV file and convert to text.

        Args:
            file_path: Path to CSV file

        Returns:
            Tuple of (text_content, metadata)
        """
        try:
            df = pd.read_csv(file_path)

            # Convert DataFrame to text (tab-separated)
            text = df.to_string(index=False)

            metadata = {
                "row_count": len(df),
                "column_count": len(df.columns),
                "columns": list(df.columns)
            }

            return text.strip(), metadata

        except Exception as e:
            logger.error(f"Error loading CSV {file_path}: {e}")
            raise

    def load_document(self, file_path: str) -> Document:
        """
        Load a document and extract its content.

        Args:
            file_path: Path to the document

        Returns:
            Document object with content and metadata
        """
        # Validate file
        is_valid, error = self.validate_file(file_path)
        if not is_valid:
            raise ValueError(f"Invalid file: {error}")

        path = Path(file_path)
        file_ext = path.suffix.lower()

        logger.info(f"Loading document: {path.name} ({file_ext})")

        # Load based on file type
        extra_metadata = {}
        if file_ext == '.pdf':
            content, extra_metadata = self.load_pdf(file_path)
        elif file_ext == '.docx':
            content, extra_metadata = self.load_docx(file_path)
        elif file_ext in ['.txt', '.md']:
            content, extra_metadata = self.load_text(file_path)
        elif file_ext == '.csv':
            content, extra_metadata = self.load_csv(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

        # Create metadata
        stat = path.stat()
        metadata = DocumentMetadata(
            file_path=str(path.absolute()),
            file_name=path.name,
            file_type=file_ext.lstrip('.'),
            file_size=stat.st_size,
            created_at=datetime.fromtimestamp(stat.st_ctime),
            modified_at=datetime.fromtimestamp(stat.st_mtime),
            indexed_at=datetime.utcnow(),
            word_count=len(content.split()),
            hash=self.compute_file_hash(file_path),
            **extra_metadata
        )

        # Generate document ID from hash
        document_id = f"doc_{metadata.hash[:16]}"

        document = Document(
            document_id=document_id,
            content=content,
            metadata=metadata
        )

        logger.info(
            f"Loaded document: {path.name} "
            f"({len(content)} chars, {metadata.word_count} words)"
        )

        return document

    def chunk_text(self, text: str, token_size: int = None) -> List[str]:
        """
        Split text into chunks with overlap.

        Args:
            text: Input text to chunk
            token_size: Approximate token count per chunk

        Returns:
            List of text chunks
        """
        if token_size is None:
            token_size = self.chunk_size

        # Simple word-based chunking (approximating tokens as words * 1.3)
        words = text.split()
        word_chunk_size = int(token_size / 1.3)
        word_overlap = int(self.chunk_overlap / 1.3)

        chunks = []
        i = 0
        while i < len(words):
            chunk_words = words[i:i + word_chunk_size]
            chunks.append(' '.join(chunk_words))
            i += word_chunk_size - word_overlap

            if i >= len(words):
                break

        return chunks

    def create_chunks(self, document: Document) -> List[DocumentChunk]:
        """
        Create chunks from a document.

        Args:
            document: Document to chunk

        Returns:
            List of DocumentChunk objects
        """
        text_chunks = self.chunk_text(document.content)

        chunks = []
        for idx, chunk_text in enumerate(text_chunks):
            chunk_id = f"{document.document_id}_chunk_{idx}"

            # Estimate token count (words * 1.3)
            word_count = len(chunk_text.split())
            token_count = int(word_count * 1.3)

            chunk = DocumentChunk(
                chunk_id=chunk_id,
                document_id=document.document_id,
                content=chunk_text,
                chunk_index=idx,
                token_count=token_count,
                metadata=document.metadata
            )
            chunks.append(chunk)

        logger.info(
            f"Created {len(chunks)} chunks for document: "
            f"{document.metadata.file_name}"
        )

        return chunks


# Global service instance
document_loader = DocumentLoaderService()
